"""DOCX 解析器：提取段落、标题层级和表格文字。"""

from pathlib import Path

from docx import Document as WordDocument

from app.parsers.base import ParsedBlock


class DocxParser:
    """使用 python-docx 把 Word 内容转换成统一 ParsedBlock。"""
    name = "python-docx"
    version = "1"

    def parse(self, path: Path) -> list[ParsedBlock]:
        document = WordDocument(path)
        blocks: list[ParsedBlock] = []
        headings: list[str] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style_name = paragraph.style.name if paragraph.style else ""
            if style_name.lower().startswith("heading"):
                try:
                    level = int(style_name.split()[-1])
                except ValueError:
                    level = 1
                headings[:] = headings[: level - 1]
                headings.append(text)
            else:
                blocks.append(ParsedBlock(content=text, section_path=headings.copy()))
        for table_number, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not rows:
                continue
            headers = rows[0]
            for row_number, values in enumerate(rows[1:] or rows, start=2 if len(rows) > 1 else 1):
                fields = [f"{headers[index] if index < len(headers) and headers[index] else f'列{index + 1}'}={value}" for index, value in enumerate(values) if value]
                if fields:
                    blocks.append(ParsedBlock(content="；".join(fields), section_path=[*headings, f"表格 {table_number}"], row_start=row_number, row_end=row_number))
        return blocks
